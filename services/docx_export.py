"""Converts the markdown report text into a .docx Document.

Handles the subset of markdown the system prompt actually produces: ATX
headings, nested bullet/numbered lists, blockquotes, pipe tables,
bold/italic/code spans, and the `[Source: "Title," Publisher, Date](https://url)`
inline citations (turned into real Word hyperlinks, same as the frontend does
for the browser view) — not a general-purpose markdown parser.

List nesting depth is derived from each line's leading indentation (any
increase in indentation opens one nesting level deeper, matching how the
frontend's markdown renderer interprets the same markdown). Ordered items are
rendered as literal "N." text with a hanging indent rather than Word's native
numPr auto-numbering: python-docx's built-in "List Number" style binds every
paragraph to the *same* numbering definition, so two unrelated numbered lists
elsewhere in the same report (e.g. the workflow map and the opportunity list)
would otherwise render as one continuously-incrementing list instead of each
starting at 1.
"""

from __future__ import annotations

import re

import docx.opc.constants
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Twips

SOURCE_CITATION_RE = re.compile(
    r'\[(Source:[^\]]+)\]\((https?://[^\s)]+)\)', re.IGNORECASE
)
INLINE_TOKEN_RE = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*(?!\*).+?(?<!\*)\*(?!\*)|`[^`]+`)")
HEADING_RE = re.compile(r"^(#{1,4})\s+(.*)")
BULLET_RE = re.compile(r"^[-*]\s+(.*)")
NUMBERED_RE = re.compile(r"^(\d+)\.\s+(.*)")
QUOTE_RE = re.compile(r"^>\s?(.*)")
RULE_RE = re.compile(r"^(-{3,}|\*{3,})$")
TABLE_SEPARATOR_RE = re.compile(r"^:?-+:?$")

LINK_COLOR = RGBColor(0x9C, 0x53, 0x26)

# Matches Word's own built-in "List Bullet"/"List Number" indent step (0.25in).
LIST_INDENT_UNIT = 360
BULLET_STYLES = ["List Bullet", "List Bullet 2", "List Bullet 3"]


def _list_depth(indent_stack: list[int], indent: int) -> int:
    """Maps a line's leading-space count to a 0-based nesting depth.

    Any indentation increase relative to the current stack opens one level
    deeper; any decrease pops back to (or below) that level — mirrors the
    frontend's `mdToHtml` nesting rule so the two renderers agree.
    """
    while indent_stack and indent < indent_stack[-1]:
        indent_stack.pop()
    if not indent_stack or indent > indent_stack[-1]:
        indent_stack.append(indent)
    return len(indent_stack) - 1


def _apply_list_indent(paragraph, depth: int) -> None:
    left = Twips(LIST_INDENT_UNIT * (depth + 1))
    paragraph.paragraph_format.left_indent = left
    paragraph.paragraph_format.first_line_indent = Twips(-LIST_INDENT_UNIT)
    return left


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), str(LINK_COLOR))
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(color)
    run_props.append(underline)
    run.append(run_props)

    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_plain_inline(paragraph, text: str) -> None:
    pos = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_inline_runs(paragraph, text: str) -> None:
    pos = 0
    for match in SOURCE_CITATION_RE.finditer(text):
        _add_plain_inline(paragraph, text[pos : match.start()])
        label, url = match.group(1), match.group(2)
        _add_hyperlink(paragraph, url, f"[{label}]")
        pos = match.end()
    _add_plain_inline(paragraph, text[pos:])


def _add_table(document: Document, table_lines: list[str]) -> None:
    rows = [[cell.strip() for cell in line.strip("|").split("|")] for line in table_lines]
    rows = [row for row in rows if not all(TABLE_SEPARATOR_RE.match(cell) for cell in row)]
    if not rows:
        return

    col_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=col_count)
    table.style = "Light Grid Accent 1"
    for row_idx, row in enumerate(rows):
        for col_idx in range(col_count):
            cell_text = row[col_idx] if col_idx < len(row) else ""
            cell = table.cell(row_idx, col_idx)
            paragraph = cell.paragraphs[0]
            _add_inline_runs(paragraph, cell_text)
            if row_idx == 0:
                for run in paragraph.runs:
                    run.bold = True
    document.add_paragraph()


def markdown_to_docx(markdown_text: str) -> Document:
    document = Document()
    lines = markdown_text.replace("\r\n", "\n").split("\n")
    i, n = 0, len(lines)
    indent_stack: list[int] = []

    while i < n:
        raw_line = lines[i].replace("\t", "    ")
        stripped = raw_line.strip()

        if not stripped:
            # Blank lines don't close a list: the report writes loose lists
            # (a blank line between each numbered step) purely for
            # readability, and it must keep counting 1, 2, 3... across them.
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            indent_stack.clear()
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _add_table(document, table_lines)
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            indent_stack.clear()
            level = len(heading_match.group(1))
            heading = document.add_heading(level=min(level, 4))
            _add_inline_runs(heading, heading_match.group(2))
            i += 1
            continue

        if RULE_RE.match(stripped):
            indent_stack.clear()
            document.add_paragraph("—" * 20)
            i += 1
            continue

        indent = len(raw_line) - len(raw_line.lstrip(" "))

        bullet_match = BULLET_RE.match(stripped)
        if bullet_match:
            depth = _list_depth(indent_stack, indent)
            style = BULLET_STYLES[min(depth, len(BULLET_STYLES) - 1)]
            paragraph = document.add_paragraph(style=style)
            _apply_list_indent(paragraph, depth)
            _add_inline_runs(paragraph, bullet_match.group(1))
            i += 1
            continue

        numbered_match = NUMBERED_RE.match(stripped)
        if numbered_match:
            depth = _list_depth(indent_stack, indent)
            paragraph = document.add_paragraph()
            left = _apply_list_indent(paragraph, depth)
            paragraph.paragraph_format.tab_stops.add_tab_stop(left, WD_TAB_ALIGNMENT.LEFT)
            paragraph.add_run(f"{numbered_match.group(1)}.\t")
            _add_inline_runs(paragraph, numbered_match.group(2))
            i += 1
            continue

        indent_stack.clear()

        quote_match = QUOTE_RE.match(stripped)
        if quote_match:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Twips(432)
            _add_inline_runs(paragraph, quote_match.group(1))
            for run in paragraph.runs:
                run.italic = True
            i += 1
            continue

        paragraph = document.add_paragraph()
        _add_inline_runs(paragraph, stripped)
        i += 1

    return document
